from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_isu_geobot_v2():
    doc = Document()
    
    # NEW RECOMMENDED TITLE
    working_title = "ISU-GeoBot: An Intelligent Geospatial Assistant with Dynamic Contextual Inference and Automated Event Summarization"

    # Title Section
    title = doc.add_heading('CAPSULE PROPOSAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Basic Info
    doc.add_paragraph("NAME: Christian Paul Simbulan & Michael Almario")
    doc.add_paragraph("PROGRAM: Bachelor of Science in Computer Science (Major in Data Mining)")
    doc.add_paragraph(f"WORKING TITLE: {working_title}")

    # Rationale
    doc.add_heading('Rationale and Significance:', level=1)
    doc.add_paragraph(
        "Current navigational aids at ISU Echague provide only static location data, leaving users unaware of active events or facility statuses. "
        "There is a need for a system that provides 'Contextual Awareness'—not just showing where a place is, but explaining what is happening there in real-time."
    )

    # Proposed Solution (UPDATED WITH CONDITIONAL SUMMARY LOGIC)
    doc.add_heading('Proposed Solution:', level=1)
    doc.add_paragraph(
        "ISU-GeoBot utilizes a RAG-powered pipeline built on Supabase to provide dynamic, situation-aware responses:"
    )
    
    layers = [
        "Dynamic Contextual Ingestion: Using Groq/Llama-3 to extract and summarize campus memos, storing them as temporal metadata in Supabase.",
        "Conditional RAG Branching: A logic layer that checks for active events at a queried location. If an event is detected, the AI generates a brief activity summary; otherwise, it defaults to providing the map pin and hyper-local landmarks.",
        "Automated Contextual Profiler (ACP): A Random Forest classifier that infers location status based on mined schedule and time-series data.",
        "Interactive Geospatial Interface: A web application that renders dynamic pins and contextual explanations via Leaflet/Folium."
    ]
    for layer in layers:
        doc.add_paragraph(layer, style='List Bullet')

    # Major Activities
    doc.add_heading('Major Activities:', level=1)
    activities = [
        "Data Mining: Building the OCR-to-JSON pipeline for event extraction and summarization.",
        "Database Architecture: Setting up Supabase with PostGIS for location indexing and pgvector for contextual search.",
        "Model Training: Developing the Random Forest status-inference model.",
        "UI/UX Development: Designing the 'Context-First' interface for dynamic map pins."
    ]
    for activity in activities:
        doc.add_paragraph(activity, style='List Number')

    doc.save("ISU_GeoBot_v2_Contextual.docx")
    print(f"Success: Proposal generated with title: {working_title}")

if __name__ == "__main__":
    create_isu_geobot_v2()